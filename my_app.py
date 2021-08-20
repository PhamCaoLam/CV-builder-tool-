from docx import Document
from docx.shared import Inches
import pyttsx3

# function used for adding a job to CV
def add_job():
    speak('In which company did you work? ')
    company = input('Company: ')
    speak('When did you start working there? ')
    from_Date = input('started working from: ')
    speak('When did you finish working there? ')
    to_Date = input('finished working on: ')
    speak('Please briefly describe your position there ')
    work_description = input('Describe briefly your work at ' + company + ': ')

    p = document.add_paragraph(company.upper() + ' : ' + from_Date + ' - ' + to_Date 
                                    + ' : ' + work_description.capitalize())
    p.style = 'List Bullet'

# function for speaking out 
# please wait until the voice finishes its text, then type your answers. 
# Otherwise, you would duplicate information. 
def speak(txt): 
    pyttsx3.speak(txt)

# create a new document
document = Document()

# profile picture 
document.add_picture('pexels.jpeg', width=Inches(0.5), height=Inches(0.5))

# name, age and origin
document.add_heading('Personal Information').bold = True

speak('What is your name?')
name = input('What is your name? ')
speak('When were you born?')
dateOfBirth = input('When were you born? ')
speak('Where are you from?')
nationality = input('Where are you from? ')


document.add_paragraph('Name : ' + name + '\n' 
                    + 'Date of birth : '+ dateOfBirth + '\n'
                    + 'Nationality : ' + nationality + '\n')

# working experience
document.add_heading('Working experience').bold = True
add_job()

# Adding more work experiences (if any) by using a while loop

while True: 
    speak('Do you have more experiences ? If yes, please type yes ?, else, type anything.')
    has_more_Experiences = input('Do you have more experiences ? ')
    if has_more_Experiences.lower() == 'yes': 
        add_job()
    else: 
        break

document.save('cv.docx')
